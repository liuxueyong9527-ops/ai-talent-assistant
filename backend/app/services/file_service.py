import uuid
import json
from pathlib import Path
from typing import Optional

from pypdf import PdfReader
from docx import Document as DocxDocument
from fastapi import UploadFile, HTTPException

from app.core.config import settings
from app.models.document import Document
from app.models.extraction_result import ExtractionResult


def ensure_upload_dir():
    settings.UPLOAD_DIR.mkdir(parents=True, exist_ok=True)


def extract_text_from_pdf(file_path: Path) -> str:
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text.strip()


def extract_text_from_docx(file_path: Path) -> str:
    doc = DocxDocument(file_path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def extract_text(file_path: Path, content_type: str) -> str:
    if "pdf" in content_type or file_path.suffix.lower() == ".pdf":
        return extract_text_from_pdf(file_path)
    elif "word" in content_type or "document" in content_type or file_path.suffix.lower() in (".docx", ".doc"):
        return extract_text_from_docx(file_path)
    raise HTTPException(status_code=400, detail="Unsupported file format. Use PDF or DOCX.")


def save_upload(
    file: UploadFile,
    user_id: int,
    doc_type: str,
) -> tuple[Document, str]:
    ensure_upload_dir()
    
    allowed = {".pdf", ".docx", ".doc"}
    suffix = Path(file.filename or "").suffix.lower()
    if suffix not in allowed:
        raise HTTPException(status_code=400, detail="Only PDF and DOC/DOCX files are allowed")

    unique_name = f"{uuid.uuid4()}{suffix}"
    file_path = settings.UPLOAD_DIR / unique_name

    content = file.file.read()
    if len(content) > settings.MAX_UPLOAD_SIZE:
        raise HTTPException(status_code=400, detail="File too large. Max 10MB.")

    with open(file_path, "wb") as f:
        f.write(content)

    raw_text = extract_text(file_path, file.content_type or "")

    relative_path = str(Path("data") / "uploads" / unique_name)
    if str(settings.UPLOAD_DIR).startswith("/"):
        relative_path = str(file_path)

    return (
        Document(
            user_id=user_id,
            type=doc_type,
            file_path=str(file_path),
            original_filename=file.filename or "unknown",
            raw_text=raw_text,
        ),
        raw_text,
    )


def parse_extraction_json(json_str: Optional[str]) -> dict:
    if not json_str:
        return {}
    try:
        return json.loads(json_str)
    except json.JSONDecodeError:
        return {}
